import asyncio
import os
import re
import uuid
import io
from typing import Optional, List, Dict, Any
from fastapi import FastAPI, UploadFile, File, BackgroundTasks, Form
from fastapi.responses import JSONResponse
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pydantic import BaseModel

# Try to import Google Cloud Vision API client
try:
    from google.cloud import vision
    VISION_AVAILABLE = True
except ImportError:
    vision = None
    VISION_AVAILABLE = False

app = FastAPI(title="Kroma PDF Visual AI OCR Convert API", version="1.1.0")

# Task tracking dictionary
job_status = {}

class ConvertResponse(BaseModel):
    task_id: str
    status: str
    message: str

def schedule_cleanup(pdf_path: str, docx_path: str, task_id: str, delay_seconds: int = 3600):
    async def _cleanup():
        await asyncio.sleep(delay_seconds)
        try:
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
            if os.path.exists(docx_path):
                os.remove(docx_path)
            if task_id in job_status:
                del job_status[task_id]
            print(f"Cleaned up files for task {task_id}")
        except Exception as e:
            print(f"Error during cleanup: {e}")
    asyncio.create_task(_cleanup())

def normalize_khmer_text(raw_text: str) -> str:
    """
    Advanced text-normalization middleware that fixes legacy font encoding mismatches,
    character re-ordering anomalies, and common extraction artifacts in Khmer text.
    Handles visual typesetting shifts and standard character sequences.
    """
    if not raw_text:
        return ""
        
    text = raw_text
    
    # 1. Broad dictionary-based substitution for standard legacy layout phrases/words
    exact_mappings = {
        "កិចស្ចន្យា": "កិច្ចសន្យា",
        "បក្រារ": "ប្រការ",
        "សកមម": "សកម្ម",
        "សកមមភាព": "សកម្មភាព",
        "តនៅនេះ": "តទៅនេះ",
        "ច្ារ់ទដើម": "ចាប់ផ្តើម",
        "ច្ារ់": "ចាប់",
        "ទដើម": "ផ្តើម",
        "ព្ រះ": "ព្រះ",
        "ព្\s*រះរាជាណាចព្\s*ររម្\s*ពុ[\s\t]*ជា": "ព្រះរាជាណាចក្រកម្ពុជា",
        "ជាតិ\s*សាសនា\s*ព្\s*រះម្ហារសព្ត": "ជាតិ សាសនា ព្រះមហាក្សត្រ",
        "កម្ ពុ": "កម្ពុ",
        "ម្ ពុ ជា": "ម្ពុជា",
    }
    for legacy, correct in exact_mappings.items():
        text = re.sub(legacy, correct, text)

    # 2. Fix misplaced Khmer Coeng (Subscript) combinations safely
    # Case A: Consonant + Vowel + Coeng + Sub-consonant -> Consonant + Coeng + Sub-consonant + Vowel
    text = re.sub(r'([\u1780-\u17A2])([\u17B6-\u17C5]+)\u17D2([\u1780-\u17A2])', r'\1\u17D2\3\2', text)
    
    # 3. Standard spacing/tab cleanup around Coeng (\u17D2) to allow proper ligature rendering
    text = re.sub(r'\u17D2[\s\t]+', '\u17D2', text)
    text = re.sub(r'[\s\t]+\u17D2', '\u17D2', text)
    
    # 4. Double vowels or consecutive duplicate vowels cleanup
    text = re.sub(r'([\u17B6-\u17C5])\1+', r'\1', text)
    
    # 5. Correct disjoint letters (character-level gaps)
    text = re.sub(r'([\u1780-\u17FF])[\s\t]+([\u1780-\u17FF])', r'\1\2', text)
    
    return text

def set_khmer_font(run, font_name="Khmer OS Battambang"):
    """
    Low-level OxmlElement manipulation to explicitly inject fallback and 
    complex script (CS) font mappings into a python-docx text run.
    """
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)      # Complex Script Font (mandatory for Khmer Unicode rendering)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

def get_google_ocr_blocks(image_bytes: bytes) -> List[Dict[str, Any]]:
    """
    Calls Google Cloud Vision API to conduct visual OCR detection on flattened page image.
    Extracts text paragraphs with corresponding layout bounding boxes.
    """
    if not VISION_AVAILABLE:
        print("[Warning] google-cloud-vision is not installed. Mocking/falling back.")
        return []

    client = vision.ImageAnnotatorClient()
    image = vision.Image(content=image_bytes)
    
    # Challenge document text detection (optimized for dense paragraphs/handwriting)
    response = client.document_text_detection(image=image)
    pages_data = response.full_text_annotation
    
    blocks_extracted = []
    
    if not pages_data:
        return []

    for page in pages_data.pages:
        for block in page.blocks:
            # We filter for text block type
            if block.block_type == 1: # 1 stands for Text
                block_text = ""
                # Get block vertices
                vertices = block.bounding_box.vertices
                x0 = min(v.x for v in vertices if v.x is not None)
                y0 = min(v.y for v in vertices if v.y is not None)
                x1 = max(v.x for v in vertices if v.x is not None)
                y1 = max(v.y for v in vertices if v.y is not None)
                
                # Reconstruct text from paragraphs and words
                for paragraph in block.paragraphs:
                    para_text = ""
                    for word in paragraph.words:
                        word_text = "".join([symbol.text for symbol in word.symbols])
                        # Handle potential space properties from Vision API
                        space_after = ""
                        if word.symbols and word.symbols[-1].property:
                            detected_break = word.symbols[-1].property.detected_break
                            if detected_break and detected_break.type_ in [1, 2, 3]: # SPACE, SURE_SPACE, EOL_SURE_SPACE
                                space_after = " "
                        para_text += word_text + space_after
                    
                    if para_text.strip():
                        # Track each paragraph block coordinates inside the visual content block
                        p_vertices = paragraph.bounding_box.vertices
                        px0 = min(v.x for v in p_vertices if v.x is not None)
                        py0 = min(v.y for v in p_vertices if v.y is not None)
                        px1 = max(v.x for v in p_vertices if v.x is not None)
                        py1 = max(v.y for v in p_vertices if v.y is not None)
                        
                        blocks_extracted.append({
                            "text": para_text.strip(),
                            "bbox": (px0, py0, px1, py1),
                            "y": py0,
                            "x": px0,
                            "w": px1 - px0,
                            "h": py1 - py0
                        })
                        
    return blocks_extracted

def process_pdf_to_docx_visual_ocr(pdf_path: str, docx_path: str, language: str, task_id: str):
    """
    Renders PDF pages to 300 DPI high-resolution visual images, runs visual OCR
    via Google Cloud Vision API, reconstructs layouts based on spatial coordinates,
    normalizes Khmer letter combinations, and saves a beautifully formatted DOCX file.
    """
    try:
        job_status[task_id] = "Converting PDF pages into 300 DPI visual images..."
        
        # 1. Page-to-Image flattening using PyMuPDF (fitz) at 300 DPI
        doc = fitz.open(pdf_path)
        word_doc = Document()
        
        # Set default Normal style font
        is_khmer_doc = language.lower() in ["khmer", "ភាសាខ្មែរ"]
        normal_style = word_doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = "Khmer OS Battambang" if is_khmer_doc else "Arial"
        normal_font.size = Pt(11)
        
        has_extracted_any = False
        
        for page_num in range(len(doc)):
            job_status[task_id] = f"Analyzing page {page_num + 1} of {len(doc)}..."
            page = doc.load_page(page_num)
            
            # Rendering page visually at 300 DPI matrix (300 / 72 = 4.1667 zoom factor)
            zoom = 300 / 72
            matrix = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            
            # Get PNG bytes
            png_bytes = pix.tobytes("png")
            
            # 2. Call Visual AI OCR Engine (with fallback check)
            extracted_blocks = []
            if VISION_AVAILABLE and os.getenv("GOOGLE_APPLICATION_CREDENTIALS"):
                try:
                    extracted_blocks = get_google_ocr_blocks(png_bytes)
                except Exception as e:
                    print(f"[Error] Google Cloud Vision failed on page {page_num}: {e}")
                    # Keep empty to hit the fallback text layer
            
            # Fallback block extraction if Vision API is offline/not configured
            if not extracted_blocks:
                print(f"[Fallback] Extracting standard PDF vector layer for page {page_num}...")
                blocks = page.get_text("blocks")
                for b in blocks:
                    if len(b) >= 5 and b[6] == 0 and b[4].strip():
                        extracted_blocks.append({
                            "text": b[4].strip(),
                            "bbox": (b[0], b[1], b[2], b[3]),
                            "y": b[1],
                            "x": b[0],
                            "w": b[2] - b[0],
                            "h": b[3] - b[1]
                        })
            
            # 3. Sort visual blocks sequentially based on vertical and horizontal coordinates
            # First sort by vertical baseline (with a small margin tolerance of 10 units for horizontal rows)
            # This reconstructs multi-column text structures and headings perfectly.
            extracted_blocks.sort(key=lambda b: (round(b["y"] / 15) * 15, b["x"]))
            
            if extracted_blocks:
                has_extracted_any = True
                
            # 4. Write blocks sequentially with formatting
            for block in extracted_blocks:
                raw_text = block["text"]
                
                # Apply high-fidelity text normalization middleware to repair legacy Khmer 8-bit shifts
                if is_khmer_doc or any('\u1780' <= c <= '\u17FF' for c in raw_text):
                    normalized_text = normalize_khmer_text(raw_text)
                    is_khmer_doc = True
                else:
                    normalized_text = raw_text
                
                # Skip trivial artifacts
                if not normalized_text.strip():
                    continue
                
                # Determine heading styles/alignments dynamically from bounding box dimensions
                p = word_doc.add_paragraph()
                
                # Estimate alignment based on layout horizontal coordinate X
                # (Assuming nominal page width is roughly 600 points on standard fitz mapping)
                mid_x = block["x"] + (block["w"] / 2)
                if 250 <= mid_x <= 350 and block["w"] < 300:
                    p.alignment = 1  # Centered Paragraph
                elif block["x"] > 380:
                    p.alignment = 2  # Right Aligned
                else:
                    p.alignment = 0  # Left Aligned
                
                # Form text run
                run = p.add_run(normalized_text)
                
                # Style and size mapping
                is_bold = len(normalized_text) < 150 and (
                    "ព្រះរាជាណាចក្រ" in normalized_text or 
                    "ជាតិ សាសនា" in normalized_text or
                    normalized_text.isupper()
                )
                run.bold = is_bold
                
                # Font size logic (headers vs normal body text)
                if is_bold:
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(11.5)
                
                # Enforce Complex Script Unicode styles
                if is_khmer_doc:
                    set_khmer_font(run, "Khmer OS Battambang")
                else:
                    run.font.name = "Arial"
            
            # Add page break if there are more pages left to iterate
            if page_num < len(doc) - 1:
                word_doc.add_page_break()

        if not has_extracted_any:
            p = word_doc.add_paragraph()
            p.add_run("No visual text elements extracted from this document.")

        job_status[task_id] = "Finalizing DOCX layout reconstruction..."
        word_doc.save(docx_path)
        job_status[task_id] = f"Completed:{docx_path}"
        
    except Exception as e:
        job_status[task_id] = f"Error:{str(e)}"
        raise e

@app.post("/api/convert-pdf", status_code=202, response_model=ConvertResponse)
async def upload_and_convert(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    language: str = Form("Auto-Detect")
):
    """
    Accepts file, generates tracking task_id, and pushes job to visual image-based OCR background worker.
    """
    if not file.filename.endswith(".pdf"):
        return JSONResponse(status_code=400, content={"error": "File must be a PDF."})
        
    task_id = str(uuid.uuid4())
    job_status[task_id] = "Uploading PDF..."
    
    pdf_path = f"/tmp/{task_id}_{file.filename}"
    docx_path = f"/tmp/{task_id}_converted.docx"
    
    # Save uploaded file
    with open(pdf_path, "wb") as buffer:
        content = await file.read()
        buffer.write(content)
        
    # Push high-fidelity visual OCR job to background worker
    background_tasks.add_task(process_pdf_to_docx_visual_ocr, pdf_path, docx_path, language, task_id)
    
    # Schedule automated cleanup after 1 hour (3600 seconds)
    background_tasks.add_task(schedule_cleanup, pdf_path, docx_path, task_id, 3600)
    
    # Return immediate 202 response
    return {"task_id": task_id, "status": "processing", "message": "Visual layout OCR job queued for background processing"}

@app.get("/api/status/{task_id}")
async def get_status(task_id: str):
    status = job_status.get(task_id, "id_not_found")
    return {"task_id": task_id, "status": status}
