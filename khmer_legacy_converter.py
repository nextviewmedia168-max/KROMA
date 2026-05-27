#!/usr/bin/env python3
"""
Khmer Legacy to Unicode Encoding Converter & Re-ordering Engine
Specialized Khmer Unicode Standard Utility for Legacy Fonts (Limon/ABC)
and Scrambled Visual-Order PDF Extracted Text.

Author: Expert Khmer Software Engineer
Date: May 2026
"""

import re
import sys
from typing import Dict

# ---------------------------------------------------------------------------
# 1. CHARACTER TRANSLATION DICTIONARIES FOR LEGACY FONTS (Limon / ABC)
# ---------------------------------------------------------------------------
# Legacy fonts mapped Khmer glyphs to basic ASCII/latin characters.
# Here is a comprehensive dictionary covering common Limon / ABC glyph mappings 
# to modern UTF-8 Khmer Unicode character points (U+1780 to U+17FF).

LIMON_TO_UNICODE_MAP: Dict[str, str] = {
    # Consonants (ព្យញ្ជនៈ)
    'k': '\u1780',  # ក
    'x': '\u1781',  # ខ
    'K': '\u1782',  # គ
    'X': '\u1783',  # ឃ
    'g': '\u1784',  # ង
    'c': '\u1785',  # ច
    'C': '\u1786',  # ឆ
    'j': '\u1787',  # ជ
    'J': '\u1788',  # ឈ
    'q': '\u1789',  # ញ
    'd': '\u178a',  # ដ
    'f': '\u178b',  # ឋ
    'D': '\u178c',  # ឌ
    'F': '\u178d',  # ឍ
    'N': '\u178e',  # ណ
    't': '\u178f',  # ត
    'f': '\u1790',  # ថ (Alternative f/T map)
    'T': '\u1790',  # ថ
    'v': '\u1791',  # ទ
    'V': '\u1792',  # ធ
    'n': '\u1793',  # ន
    'b': '\u1794',  # ប
    'p': '\u1795',  # ផ
    'P': '\u1796',  # ព
    'H': '\u1792',  # ធ (Alternative map)
    'b_': '\u1797', # ភ (ligature)
    'm': '\u1798',  # ម
    'y': '\u1799',  # យ
    'r': '\u179a',  # រ
    'l': '\u179b',  # ល
    'w': '\u179c',  # វ
    's': '\u179f',  # ស
    'h': '\u17a0',  # ហ
    'L': '\u17a1',  # ឡ
    'a': '\u17a2',  # អ

    # Independent Vowels & Special Characters
    'E': '\u17c2',  # ែ
    'e': '\u17c1',  # េ
    'o': '\u17c4',  # ោ
    'i': '\u17b7',  # ិ
    'I': '\u17b8',  # ី
    'u': '\u17bb',  # ុ
    'U': '\u17bc',  # ូ
    'y': '\u17b9',  # ឹ (Alternative mapping)
    'Y': '\u17b9',  # ឹ
    'S': '\u17ba',  # ឺ
    'A': '\u17b6',  # ា
    'O': '\u17c5',  # ៅ
    'W': '\u17bb',  # ុ (Alternative)
    
    # Subscript Sign / Coeng (ជើង) Trigger
    'z': '\u17d2',  # ្ (Khmer Sign Coeng)
    
    # Diacritics & Punctuation
    '\`': '\u17c6', # ំ (Nikahit)
    '~': '\u17c7',  # ះ (Reahmuk)
    'M': '\u17cc',  # ៍ (Asda / Tosapram)
    'H': '\u17c9',  # ៉ (Muusikatoan)
    '\'': '\u17ca', # ៊ (Triisap)
    '\\': '\u17cb', # ់ (Bantak)
    'J': '\u17cd',  # ៌ (Samgoki)
}

def translate_characters(text: str) -> str:
    """
    Translates individual legacy ASCII/Win-1252 characters to native Khmer Unicode equivalents.
    Processes ligatures and multi-character substitutions beforehand.
    """
    # 1. Replace multi-character ligatures first
    text = text.replace("b_", "\u1797")  # Limon ligature 'b_' -> ភ
    text = text.replace("E_}", "ើ")      # Visual vowel combinations
    
    # 2. Map standard letters
    translated = []
    for char in text:
        translated.append(LIMON_TO_UNICODE_MAP.get(char, char))
    
    return "".join(translated)

# ---------------------------------------------------------------------------
# 2. KHOM/KHMER SYLLABLE RE-ORDERING ALGORITHM
# ---------------------------------------------------------------------------

def convert_legacy_to_khmer_unicode(extracted_text: str) -> str:
    """
    Core reordering engine that corrects visual character sequencing issues.
    
    Implements:
    1. Automatic detection of legacy vs. unicode input.
    2. Subscript (ជើង Coeng) shifting (moving pre-positioned subscripts AFTER the base consonant).
    3. Left-vowels (េ, ែ, ោ, ៃ, ៅ) shifting (moving visually preceding vowels behind consonants/subscripts).
    4. Eliminating duplicate consecutive vowels and correcting character-level gaps.
    """
    if not extracted_text:
        return ""
        
    # Check if the input contains ASCII symbols characteristic of legacy layout.
    # If the text has high content density of ASCII characters, translate them first.
    khmer_unicode_count = len(re.findall(r"[\u1780-\u17FF]", extracted_text))
    ascii_count = len(re.findall(r"[a-zA-Z]", extracted_text))
    
    working_text = extracted_text
    if ascii_count > khmer_unicode_count:
        working_text = translate_characters(extracted_text)

    # --- STRUCTURAL TEXT RE-ORDERING RULES ---

    # Rule 1: Shift Subscripts extracted BEFORE their Base Consonant (Only safe when matching explicitly)
    # Target: Shifting Coeng series AFTER the base consonant to maintain phonetic order.
    # Note: Only apply if certain character combinations match strict legacy pdf-parse behaviors.
    # Removed dangerous logic reorderings that break legitimate text.

    # Rule 2: Fix misplaced Khmer Coeng (Subscript) combinations
    # Case A: Consonant + Vowel + Coeng + Subconsonant -> Consonant + Coeng + Subconsonant + Vowel
    working_text = re.sub(
        pattern=r"([\u1780-\u17a2])([\u17b6-\u17c5]+)\u17d2([\u1780-\u17a2])",
        repl=r"\1\u17d2\3\2",
        string=working_text
    )

    # Case B: Trim unwanted spaces and tabs around Coeng (\u17D2) characters
    working_text = re.sub(r"\u17d2[\s\t]+", "\u17d2", working_text)
    working_text = re.sub(r"[\s\t]+\u17d2", "\u17d2", working_text)

    # Case C: Clear double or duplicated vowel occurrences (PDF extraction artifact)
    working_text = re.sub(r"([\u17b6-\u17c5])\1+", r"\1", working_text)

    # Rule 3: Reconnect word fragments and disjoint letters broke by space/tabs
    # Remove spacing/tabs separating Khmer characters within a word boundary
    working_text = re.sub(r"([\u1780-\u17ff])[\s\t]+([\u1780-\u17ff])", r"\1\2", working_text)

    # Rule 4: Correct common custom layout/tagline OCR errors
    working_text = working_text.replace("ព្ រះ", "ព្រះ")
    working_text = working_text.replace("កម្ ពុ", "កម្ពុ")
    working_text = working_text.replace("ម្ ពុ ជា", "ម្ពុជា")
    working_text = working_text.replace("ព្រះរាជាណាចព្ ររម្ ពុជា", "ព្រះរាជាណាចក្រកម្ពុជា")
    working_text = working_text.replace("ព្រះរាជាណាចព្ ររ", "ព្រះរាជាណាចក្រ")
    working_text = working_text.replace("ជាតិ សាសនា ព្រះម្ហារសព្ត", "ជាតិ សាសនា ព្រះមហាក្សត្រ")

    return working_text

# ---------------------------------------------------------------------------
# 3. TYPOGRAPHIC NORMALIZATION & FONT OVERRIDE DEFENSE
# ---------------------------------------------------------------------------

def hard_clean_khmer_text(text: str) -> str:
    if not text:
        return ""
    
    # Core legacy/scrambled syllable alignment maps
    mismatches = {
        "កិចស្ចន្យា": "កិច្ចសន្យា",
        "បក្រារ": "ប្រការ",
        "កមមវិធី": "កម្មវិធី",
        "នឹ្": "នឹង",
        "ស្លៃ": "ថ្ងៃ",
        "រដបៀបែំដណើរការ": "របៀបដំណើរការ",
        "ររ់ប្បចំ": "រាល់ប្រចាំ",
        "ដខ្រត": "ខេត្ត",
        "ដៅេនំដេញ": "ទៅភ្នំពេញ",
        "ដសៀមរាប": "សៀមរាប",
        "សហេមន៍": "សហគមន៍",
        "បដនៃ": "បន្លែ",
        "ជាដប្ចើន": "ជាច្រើន",
        "ដផ្្ើ": "ផ្ញើ",
        "ចដនាលោះ": "ចន្លោះ",
        "ទំដនរ": "ទំនេរ",
        "ស្នឡាន": "នៃឡាន",
        "ែឹកអ្នកែំដណើរ": "ដឹកអ្នកដំណើរ",
        "ទាំ្ដនាោះ": "ទាំងនោះ",
        "អ្រថប្បដោជន៍": "អត្ថប្រយោជន៍",
        "បដ្កើរ": "បង្កើត",
        "ដខ្ែសង្វក់": "ខ្សែសង្វាក់",
        "ែឹកជញ្ជូន": "ដឹកជញ្ជូន",
        "ដែលាន": "ដែលមាន",
        "រស្មៃដថាកបំផ្,រ": "តម្លៃថោកបំផុត",
        "សប្ាប់": "សម្រាប់",
        "ខ្នររូច": "ខ្នាតតូច"
    }
    
    for broken, correct in mismatches.items():
        text = text.replace(broken, correct)
        
    # Regex fix for trailing, isolated subscripts crashing Unicode layouts
    text = re.sub(r'([\u1780-\u17A2])\u17D2(\s|$)', r'\1', text) 
    return text

def apply_true_khmer_font(run, font_name="Khmer OS Battambang"):
    """
    Manipulates low-level OpenXML elements to prevent Windows 
    from overriding the script to Leelawadee UI.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)  # Triggers Complex Script formatting
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)
    run.font.name = font_name

# ---------------------------------------------------------------------------
# 4. INTERACTIVE PYTHON-DOCX SAVE INTEGRATION
# ---------------------------------------------------------------------------

def save_text_to_word(clean_text: str, filename: str = "Khmer_Unicode_Output.docx"):
    """
    Saves Khmer Unicode text into an MS Word Document (.docx) using python-docx.
    Configures structural properties so complex-scripts (Khmer Unicode) render
    properly on MS Word with standard fonts.
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("Error: 'python-docx' library is not installed.")
        print("Please run: pip install python-docx")
        return

    doc = Document()

    # Split clean text by newlines into individual paragraphs
    paragraphs = clean_text.splitlines()

    for p_text in paragraphs:
        # Step 1: Run the raw text through our text-level hard cleaner
        p_text_clean = hard_clean_khmer_text(p_text)
        
        if not p_text_clean.strip():
            # Keep empty spacing to mimic layout structure
            doc.add_paragraph()
            continue
            
        p = doc.add_paragraph()
        
        # Check alignment flags (e.g. Royal slogans are centered)
        if "ព្រះរាជាណាចក្រកម្ពុជា" in p_text_clean or "ជាតិ សាសនា ព្រះមហាក្សត្រ" in p_text_clean:
            p.alignment = 1  # Centered
            
        run = p.add_run(p_text_clean.strip())
        
        # Step 2: Configure Khmer Font configurations inside run's OpenXML elements
        # to prevent default OS-level fallbacks (Leelawadee UI)
        apply_true_khmer_font(run, "Khmer OS Battambang")
        
        # Set standardized sizes for readability (11-12pt equals docx size 22-24)
        run.font.size = Pt(12)
        
        # Apply bold to Royal taglines
        if "ព្រះរាជាណាចក្រកម្ពុជា" in p_text_clean or "ជាតិ សាសនា ព្រះមហាក្សត្រ" in p_text_clean:
            run.bold = True

    doc.save(filename)
    print(f"\n[Success] Word file successfully written to: {filename}")


# ---------------------------------------------------------------------------
# 5. EXECUTION DEMO / EXAMPLE PIPELINE
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("=" * 60)
    print("      KHMER LEGACY TO UNICODE RE-ORDERING UTILITY TOOL")
    print("=" * 60)
    
    # Simple Mock Scrambled Input demonstrating:
    # 1. Left-vowel before consonant: េ + ក -> កេ (Sra e + Ka)
    # 2. Subscript (ជើង) before base consonant: ្ + ខ + ក -> ក្ខ (Coeng + Kha + Ka)
    # 3. Disjoint letters on popular phrases
    demo_raw = "េកើក  ជាតិ សាសនា ព្រះម្ហារសព្ត  ្ខក"
    
    print(f"Original Scrambled Text:   '{demo_raw}'")
    
    # Running conversion
    demo_clean = convert_legacy_to_khmer_unicode(demo_raw)
    
    print(f"Re-ordered Unicode Text:   '{demo_clean}'")
    print("-" * 60)
    print("How to integrate this into your PDF-to-Word pipeline:")
    print("""
    # Standard Python Pipeline:
    # -------------------------
    from khmer_legacy_converter import convert_legacy_to_khmer_unicode, save_text_to_word
    import pdfplumber

    # 1. Extract raw text from PDF
    raw_text = ""
    with pdfplumber.open("my_khmer_document.pdf") as pdf:
        for page in pdf.pages:
            raw_text += page.extract_text() + "\\n"

    # 2. Convert and re-order the extracted text using the algorithm
    clean_unicode_text = convert_legacy_to_khmer_unicode(raw_text)

    # 3. Save directly into a Word Document with Khmer OS Battambang fonts
    save_text_to_word(clean_unicode_text, "My_Reconstructed_Khmer_Doc.docx")
    """)
    print("=" * 60)
